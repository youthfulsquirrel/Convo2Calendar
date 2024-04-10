import streamlit as st
import pandas as pd
import nltk
from nltk import word_tokenize, pos_tag, FreqDist, trigrams
from nltk.corpus import stopwords
from wordcloud import WordCloud
from textblob import TextBlob
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Cm
import io
import base64
import traceback

# Download NLTK resources
nltk.download('punkt')
nltk.download('stopwords')
doc_path = "../runs/Part2/run1/output.docx"

# Function to analyze data and generate Word document
def analyze_data(csv_file, transcript_text):
    try:
        # Read CSV file
        df = pd.read_csv(csv_file)
        
        # Calculate speaker durations
        speaker_totals = df.groupby('speaker_id')['duration'].sum().sort_values(ascending=False)
        max_speaker = speaker_totals.idxmax()
        min_speaker = speaker_totals.idxmin()
        max_duration = speaker_totals[max_speaker]
        min_duration = speaker_totals[min_speaker]
        average_duration = speaker_totals.mean()

        # Create a Word document
        doc = Document()
        doc.add_heading('Meeting Analysis', level=1)
        doc.add_heading('Speaker Analysis', level=3)
        doc.add_paragraph(f"Spoke the most: {max_speaker} ({max_duration:.2f} seconds)")
        doc.add_paragraph(f"Spoke the least: {min_speaker} ({min_duration:.2f} seconds)")
        doc.add_paragraph(f"Average speaking time: {average_duration:.2f} seconds")

        #create meta data
        meta_data = {'max_speaker_id': max_speaker, 'max_speaker_duration':max_duration.round(2), 'min_speaker_id': min_speaker, 'min_speaker':min_duration.round(2), "average_duration":average_duration.round(2), 'total_speaker': len(speaker_totals)}
       
        # Tokenize transcript and perform NLP tasks
        sentences = nltk.sent_tokenize(transcript_text)
        filtered_tokens = [[word.lower() for word in word_tokenize(sentence) if word.lower() not in stopwords.words('english') and word.lower().isalnum()] for sentence in sentences]
        all_words = [word for sentence_words in filtered_tokens for word in sentence_words]
        word_freq = FreqDist(all_words)
        top_words = word_freq.most_common(10)

        total_words = len(all_words)
        meta_data['total_words'] = total_words

        # Common trigrams
        trigram_freq = FreqDist(nltk.trigrams(all_words))
        common_trigrams = trigram_freq.most_common(5)
        doc.add_heading('Common Triagrams', 3)
        doc.add_paragraph(f"1st common trigram:{common_trigrams[0][0] }, (Count:{common_trigrams[0][1]} )")
        doc.add_paragraph(f"2nd common trigram:{common_trigrams[4][0] }, (Count:{common_trigrams[4][1]} )")
        trigram_dict ={'trigram_list':[common_trigrams[0][0], common_trigrams[4][0]], 'count_list': [common_trigrams[0][1],common_trigrams[4][1]]}
        meta_data['trigrams'] = trigram_dict

        # Sentiment Analysis
        sentiments = [TextBlob(sentence).sentiment.polarity for sentence in sentences]
        positivity_score = sum(1 for sentiment in sentiments if sentiment > 0) / len(sentiments)
        positivity = f"{positivity_score:.3f}"
        negativity_score = sum(1 for sentiment in sentiments if sentiment < 0) / len(sentiments)
        negativity = f"{negativity_score:.3f}"
        neutrality_score = sum(1 for sentiment in sentiments if sentiment == 0) / len(sentiments)
        neutrality = f"{neutrality_score:.3f}"
        doc.add_heading('Sentiment Analysis', 3)
        doc.add_paragraph(f"Positivity Score: {positivity}")
        doc.add_paragraph(f"Neutrality Score: {neutrality}")
        doc.add_paragraph(f"Negativity Score:{negativity}")
        meta_data['sentiments'] = [positivity,neutrality,negativity]

        # Plot top 10 meaningful words frequency
        plt.bar(range(len(top_words)), [freq for word, freq in top_words], align='center')
        plt.xticks(range(len(top_words)), [word for word, freq in top_words])
        plt.xlabel('Words')
        plt.ylabel('Frequency')
        plt.title('Top 10 meaningful words')
        top_words_path = '../runs/Part2/run1/top_words.png'
        plt.savefig(top_words_path)
        plt.close()

        # Word cloud generation
        wordcloud = WordCloud().generate_from_frequencies(word_freq)
        plt.figure(figsize=(10,5))
        plt.imshow(wordcloud, interpolation = "bilinear")
        plt.axis('off')
        word_cloud_path = f'../runs/Part2/run1/word_cloud.png'
        plt.savefig(word_cloud_path, dpi = 200)
        
        doc.add_page_break() 
        doc.add_heading('Word Frequency',3)
        doc.add_paragraph(f'Total Word Count: {total_words}')
        doc.add_picture(top_words_path, height = Cm(12) )
        doc.add_heading('Word Cloud',3)
        doc.add_picture(word_cloud_path, width = Cm(15))
        
        with open(doc_path, "wb") as f:
            doc.save(f)
        return meta_data

    except Exception as e:
        traceback.print_exc()
        st.text(f"An error occurred: {str(e)}")
        print('+++++++++++++++++++++++++++++++++++++++__++++++++++++++++++++++++++++++')

# Streamlit app
def main():
    st.title("Meeting Analysis")

    # File uploader for CSV file
    csv_file = st.file_uploader("Upload CSV file", type=["csv"])

    # Text input for transcript
    transcript_text = st.text_area("Enter Transcript")

    # Perform analysis on button click
    if st.button("Analyze"):
        if csv_file is not None and transcript_text:
            meta_data = analyze_data(csv_file, transcript_text)
           
            # Display analysis results
            st.header("Results")
            # Provide download link for Word document
            download_link = generate_download_link(doc_path, "Download Word Document")
            st.markdown(download_link, unsafe_allow_html=True)
            st.write('**JSON**')
            st.markdown(meta_data)
        elif csv_file is None:
            st.text('Please upload csv file to proceed')
        elif transcript_text is None:
            st.text('Transcript cannot be empty')

def generate_download_link(file_path, link_text):
    with open(file_path, "rb") as file:
        data = file.read()
        b64 = base64.b64encode(data).decode('utf-8')
        href = f'<a href="data:application/octet-stream;base64,{b64}" download="run1_output.docx">{link_text}</a>'
    return href

if __name__ == "__main__":
    main()